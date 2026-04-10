#!/usr/bin/env python3
import argparse
import json
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Pt
from docx.text.paragraph import Paragraph


def load_manifest(path: Path) -> dict[str, Any]:
    suffix = path.suffix.lower()
    if suffix == '.json':
        return json.loads(path.read_text(encoding='utf-8'))
    if suffix in {'.yaml', '.yml'}:
        try:
            import yaml  # type: ignore
        except ImportError as exc:
            raise SystemExit('YAML manifest requires PyYAML; use JSON or install pyyaml.') from exc
        return yaml.safe_load(path.read_text(encoding='utf-8'))
    raise SystemExit('Manifest must be .json, .yaml, or .yml')


def insert_paragraph_after(paragraph: Paragraph, text: str = '', style: str | None = None) -> Paragraph:
    new_p = OxmlElement('w:p')
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def normalize_text(text: str) -> str:
    return ' '.join(text.split())


def find_anchor_paragraph(doc: Document, chapter: str | None, anchor_text: str | None) -> Paragraph:
    candidates = []
    if anchor_text:
        candidates.append(normalize_text(anchor_text))
    if chapter:
        candidates.append(normalize_text(chapter))
    if not candidates:
        raise SystemExit('Each item needs chapter or anchor_text')
    for para in doc.paragraphs:
        value = normalize_text(para.text)
        for candidate in candidates:
            if value == candidate or candidate in value:
                return para
    raise SystemExit(f'Anchor not found for {anchor_text or chapter}')


def set_zero_spacing(paragraph: Paragraph) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)


def configure_intro_or_analysis(paragraph: Paragraph) -> None:
    set_zero_spacing(paragraph)


def configure_figure_paragraph(paragraph: Paragraph) -> None:
    fmt = paragraph.paragraph_format
    fmt.line_spacing = 1.5
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def configure_caption_paragraph(paragraph: Paragraph) -> None:
    set_zero_spacing(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_asset_block(doc: Document, item: dict[str, Any]) -> None:
    anchor = find_anchor_paragraph(doc, item.get('chapter'), item.get('anchor_text'))
    current = anchor

    intro_text = (item.get('intro_text') or '').strip()
    if intro_text:
        current = insert_paragraph_after(current, intro_text)
        configure_intro_or_analysis(current)

    current = insert_paragraph_after(current)
    asset_path = Path(item['asset_path']).expanduser().resolve()
    if not asset_path.exists():
        raise SystemExit(f'Asset not found: {asset_path}')
    run = current.add_run()
    run.add_picture(str(asset_path))
    configure_figure_paragraph(current)

    caption = (item.get('caption') or '').strip()
    figure_number = (item.get('figure_number') or '').strip()
    caption_text = caption
    if figure_number and caption:
        caption_text = f'{figure_number} {caption}'
    elif figure_number:
        caption_text = figure_number
    if caption_text:
        current = insert_paragraph_after(current, caption_text)
        configure_caption_paragraph(current)

    analysis_text = (item.get('analysis_text') or '').strip()
    if analysis_text:
        current = insert_paragraph_after(current, analysis_text)
        configure_intro_or_analysis(current)


def main() -> None:
    parser = argparse.ArgumentParser(description='Batch insert drawio figures, UI screenshots, and code screenshots into a thesis docx.')
    parser.add_argument('manifest', help='Path to JSON/YAML asset manifest')
    parser.add_argument('--save-as', help='Output docx path; defaults to overwrite when --in-place is set')
    parser.add_argument('--in-place', action='store_true', help='Overwrite source docx')
    args = parser.parse_args()

    manifest_path = Path(args.manifest).expanduser().resolve()
    data = load_manifest(manifest_path)
    docx_path = Path(data['docx_path']).expanduser().resolve()
    if not docx_path.exists():
        raise SystemExit(f'DOCX not found: {docx_path}')
    items = data.get('items') or []
    if not items:
        raise SystemExit('Manifest items cannot be empty')

    doc = Document(str(docx_path))
    for item in items:
        add_asset_block(doc, item)

    if args.save_as:
        output_path = Path(args.save_as).expanduser().resolve()
    elif args.in_place:
        output_path = docx_path
    else:
        output_path = docx_path.with_name(docx_path.stem + '.assets.docx')
    doc.save(str(output_path))
    print(f'Wrote {output_path}')


if __name__ == '__main__':
    main()
